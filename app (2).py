# -*- coding: utf-8 -*-
"""
================================================================================
ENTERPRISE KNOWLEDGE MANAGEMENT & STRATEGIC LESSONS LEARNED
================================================================================
Desain: Apple Design System (Bento Box, Macro Typography, Monochromatic)
Tanpa emoji di seluruh antarmuka. Persistensi data via SQLite.
================================================================================
"""

import streamlit as st
import sqlite3
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
from datetime import datetime, date
import io
import os
import re
import html as html_lib

# ------------------------------------------------------------------------------
# IMPORT OPSIONAL UNTUK PARSING DOKUMEN (graceful fallback, tidak boleh crash)
# ------------------------------------------------------------------------------
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ==============================================================================
# 1. KONFIGURASI GLOBAL & KONSTANTA
# ==============================================================================

DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "km_platform.db")

STATUS_DRAFT = "Draft"
STATUS_VERIFIED = "Verified"
STATUS_REJECTED = "Rejected"
ALL_STATUSES = [STATUS_DRAFT, STATUS_VERIFIED, STATUS_REJECTED]

IMPACT_LEVELS = ["Tinggi", "Sedang", "Rendah"]

CATEGORY_OPTIONS = [
    "Perencanaan Proyek", "Manajemen Risiko", "Pengadaan dan Kontrak",
    "Kualitas dan Kepatuhan", "Sumber Daya Manusia", "Teknologi dan Sistem",
    "Keuangan dan Anggaran", "Operasional", "Stakeholder dan Komunikasi", "Lainnya"
]

# Kata kunci mesin ekstraksi otomatis "Smart Auto-Fill"
KEYWORDS_SUMMARY = [
    "isu", "masalah", "kendala", "permasalahan", "issue", "problem",
    "hambatan", "tantangan", "ditemukan bahwa", "terjadi"
]
KEYWORDS_ROOT_CAUSE = [
    "akar masalah", "akar penyebab", "disebabkan", "root cause", "penyebab",
    "faktor penyebab", "diakibatkan", "dikarenakan", "sumber masalah", "berawal dari"
]
KEYWORDS_RECOMMENDATION = [
    "rekomendasi", "solusi", "usulan", "tindak lanjut", "saran", "recommendation",
    "action plan", "mitigasi", "langkah perbaikan", "perlu dilakukan", "disarankan"
]

# ------------------------------------------------------------------------------
# TOKEN WARNA - APPLE DESIGN SYSTEM (High-Contrast Monochromatic)
# ------------------------------------------------------------------------------
COLOR_BG = "#F5F5F7"            # Latar aplikasi: abu-abu premium Apple
COLOR_CARD = "#FFFFFF"          # Latar Bento Box: putih bersih
COLOR_TEXT_PRIMARY = "#000000"  # Teks utama: hitam pekat
COLOR_TEXT_SECONDARY = "#86868B" # Teks sekunder: abu-abu elegan khas Apple
COLOR_BORDER = "rgba(0,0,0,0.03)"  # Border sangat tipis, nyaris tak terlihat

COLOR_ACCENT = "#0071E3"        # Satu-satunya aksen kuat: Apple Blue
COLOR_ACCENT_DARK = "#0058B0"

# Status/badge: warna pucat/desaturasi, bukan neon
STATUS_COLORS = {
    STATUS_VERIFIED: {"bg": "#E3F7EA", "text": "#1E7D42"},
    STATUS_DRAFT:    {"bg": "#FFF6E1", "text": "#9A6400"},
    STATUS_REJECTED: {"bg": "#FDECEC", "text": "#B23B3B"},
}
IMPACT_COLORS = {
    "Tinggi": {"bg": "#FDECEC", "text": "#B23B3B"},
    "Sedang": {"bg": "#FFF6E1", "text": "#9A6400"},
    "Rendah": {"bg": "#E3F7EA", "text": "#1E7D42"},
}
COLOR_HIGHLIGHT = "#FFE39B"  # Highlight teks hasil pencarian

_APPLE_FONT_STACK = ("-apple-system, BlinkMacSystemFont, 'SF Pro Display', "
                     "'SF Pro Text', 'Inter', 'Helvetica Neue', Arial, sans-serif")

# ------------------------------------------------------------------------------
# TEMA PLOTLY GLOBAL - monokrom/minimalis, latar transparan menyatu dengan Bento
# ------------------------------------------------------------------------------
try:
    _apple_template = pio.templates["plotly_white"]
    _apple_template.layout.font = dict(family=_APPLE_FONT_STACK, color=COLOR_TEXT_PRIMARY, size=13)
    _apple_template.layout.paper_bgcolor = "rgba(0,0,0,0)"
    _apple_template.layout.plot_bgcolor = "rgba(0,0,0,0)"
    pio.templates["apple_mono"] = _apple_template
    pio.templates.default = "apple_mono"
except Exception:
    pass


# ==============================================================================
# 2. LAPISAN DATABASE (SQLite) - PERSISTENSI DATA
# ==============================================================================

def get_connection():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    """Inisialisasi tabel database jika belum ada. Aman dipanggil berulang kali."""
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS issues (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                project_name TEXT,
                category TEXT,
                impact_level TEXT,
                status TEXT DEFAULT 'Draft',
                summary TEXT,
                root_cause TEXT,
                recommendation TEXT,
                uploader TEXT,
                upload_date TEXT,
                project_year INTEGER,
                document_version TEXT,
                file_name TEXT,
                extracted_text TEXT,
                reviewer_notes TEXT,
                reviewed_by TEXT,
                reviewed_date TEXT
            )
        """)
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Gagal inisialisasi database: {e}")


def insert_issue(data: dict) -> bool:
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO issues (
                title, project_name, category, impact_level, status,
                summary, root_cause, recommendation, uploader, upload_date,
                project_year, document_version, file_name, extracted_text,
                reviewer_notes, reviewed_by, reviewed_date
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            data.get("title", "").strip(),
            data.get("project_name", "").strip(),
            data.get("category", "Lainnya"),
            data.get("impact_level", "Sedang"),
            data.get("status", STATUS_DRAFT),
            data.get("summary", "").strip(),
            data.get("root_cause", "").strip(),
            data.get("recommendation", "").strip(),
            data.get("uploader", "Anonim").strip(),
            data.get("upload_date", datetime.now().strftime("%Y-%m-%d %H:%M")),
            data.get("project_year", datetime.now().year),
            data.get("document_version", "v1.0"),
            data.get("file_name", ""),
            data.get("extracted_text", ""),
            data.get("reviewer_notes", ""),
            data.get("reviewed_by", ""),
            data.get("reviewed_date", ""),
        ))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Gagal menyimpan data ke database: {e}")
        return False


def update_issue_status(issue_id: int, new_status: str, reviewer_notes: str = "", reviewed_by: str = "PMO/Manager") -> bool:
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""
            UPDATE issues
            SET status = ?, reviewer_notes = ?, reviewed_by = ?, reviewed_date = ?
            WHERE id = ?
        """, (new_status, reviewer_notes, reviewed_by,
              datetime.now().strftime("%Y-%m-%d %H:%M"), issue_id))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Gagal memperbarui status: {e}")
        return False


def delete_issue(issue_id: int) -> bool:
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM issues WHERE id = ?", (issue_id,))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Gagal menghapus data: {e}")
        return False


def fetch_all_issues() -> pd.DataFrame:
    try:
        conn = get_connection()
        df = pd.read_sql_query("SELECT * FROM issues ORDER BY id DESC", conn)
        conn.close()
        return df
    except Exception as e:
        st.error(f"Gagal mengambil data dari database: {e}")
        return pd.DataFrame(columns=[
            "id", "title", "project_name", "category", "impact_level", "status",
            "summary", "root_cause", "recommendation", "uploader", "upload_date",
            "project_year", "document_version", "file_name", "extracted_text",
            "reviewer_notes", "reviewed_by", "reviewed_date"
        ])


# ==============================================================================
# 3. MODUL PARSING DOKUMEN (PDF / DOCX / TXT)
# ==============================================================================

def extract_text_from_pdf(uploaded_file) -> str:
    raw_bytes = uploaded_file.read()
    text_parts = []
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            if text_parts:
                return "\n".join(text_parts)
        except Exception:
            text_parts = []

    if PYPDF_AVAILABLE:
        try:
            reader = PdfReader(io.BytesIO(raw_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as e:
            st.warning(f"File PDF tampak rusak atau terenkripsi, tidak dapat diekstrak sepenuhnya ({e}).")
            return ""

    st.warning("Library pembaca PDF tidak tersedia di lingkungan ini.")
    return ""


def extract_text_from_docx(uploaded_file) -> str:
    if not DOCX_AVAILABLE:
        st.warning("Library python-docx tidak tersedia di lingkungan ini.")
        return ""
    try:
        raw_bytes = uploaded_file.read()
        document = docx.Document(io.BytesIO(raw_bytes))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        st.warning(f"File DOCX tampak rusak atau tidak valid ({e}).")
        return ""


def extract_text_from_txt(uploaded_file) -> str:
    try:
        raw_bytes = uploaded_file.read()
        try:
            return raw_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return raw_bytes.decode("latin-1", errors="ignore")
    except Exception as e:
        st.warning(f"Gagal membaca file TXT ({e}).")
        return ""


def parse_uploaded_document(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    try:
        filename = uploaded_file.name.lower()
        if filename.endswith(".pdf"):
            return extract_text_from_pdf(uploaded_file)
        elif filename.endswith(".docx"):
            return extract_text_from_docx(uploaded_file)
        elif filename.endswith(".txt"):
            return extract_text_from_txt(uploaded_file)
        else:
            st.warning("Format file tidak didukung. Gunakan PDF, DOCX, atau TXT.")
            return ""
    except Exception as e:
        st.error(f"Terjadi kesalahan tak terduga saat memproses dokumen: {e}")
        return ""


# ==============================================================================
# 4. MESIN "SMART AUTO-FILL" - EKSTRAKSI KALIMAT KUNCI
# ==============================================================================

def split_into_sentences(text: str) -> list:
    if not text or not text.strip():
        return []
    cleaned = re.sub(r"\s+", " ", text).strip()
    sentences = re.split(r"(?<=[.!?])\s+", cleaned)
    return [s.strip() for s in sentences if len(s.strip()) > 15]


def extract_relevant_sentences(sentences: list, keywords: list, max_sentences: int = 3) -> str:
    matched = []
    for sentence in sentences:
        lower_sentence = sentence.lower()
        if any(kw in lower_sentence for kw in keywords):
            matched.append(sentence)
        if len(matched) >= max_sentences:
            break
    return " ".join(matched)


def smart_auto_fill(extracted_text: str) -> dict:
    """Menganalisis teks dokumen dan menyarankan isi Ringkasan, Akar Masalah,
    dan Rekomendasi berdasarkan pola kata kunci (heuristik, berjalan offline)."""
    result = {"summary": "", "root_cause": "", "recommendation": ""}
    if not extracted_text or not extracted_text.strip():
        return result
    try:
        sentences = split_into_sentences(extracted_text)
        if not sentences:
            return result

        summary = extract_relevant_sentences(sentences, KEYWORDS_SUMMARY, max_sentences=3)
        root_cause = extract_relevant_sentences(sentences, KEYWORDS_ROOT_CAUSE, max_sentences=3)
        recommendation = extract_relevant_sentences(sentences, KEYWORDS_RECOMMENDATION, max_sentences=3)

        if not summary:
            summary = " ".join(sentences[:2])
        if not root_cause:
            root_cause = "Pola akar masalah tidak ditemukan secara eksplisit. Lengkapi secara manual."
        if not recommendation:
            recommendation = "Pola rekomendasi tidak ditemukan secara eksplisit. Lengkapi secara manual."

        result["summary"] = summary
        result["root_cause"] = root_cause
        result["recommendation"] = recommendation
        return result
    except Exception as e:
        st.warning(f"Smart Auto-Fill tidak dapat memproses teks sepenuhnya ({e}).")
        return result


# ==============================================================================
# 5. UTILITAS: HIGHLIGHTING, EXPORT
# ==============================================================================

def highlight_keyword(text: str, keyword: str) -> str:
    if not text:
        return ""
    safe_text = html_lib.escape(str(text))
    if not keyword or not keyword.strip():
        return safe_text
    try:
        safe_keyword = html_lib.escape(keyword.strip())
        pattern = re.compile(re.escape(safe_keyword), re.IGNORECASE)
        highlighted = pattern.sub(
            lambda m: f'<mark style="background-color:{COLOR_HIGHLIGHT};color:{COLOR_TEXT_PRIMARY};padding:0 3px;border-radius:4px;">{m.group(0)}</mark>',
            safe_text
        )
        return highlighted
    except Exception:
        return safe_text


def dataframe_to_excel_bytes(df: pd.DataFrame) -> bytes:
    output = io.BytesIO()
    try:
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Lessons_Learned")
        return output.getvalue()
    except Exception as e:
        st.error(f"Gagal membuat file Excel: {e}")
        return b""


# ==============================================================================
# 6. CUSTOM CSS - APPLE DESIGN SYSTEM (BENTO BOX, MACRO TYPOGRAPHY)
# ==============================================================================

def load_custom_css():
    """Menyuntikkan seluruh gaya visual: tipografi raksasa, latar monokrom,
    Bento Box dengan kedalaman Z-axis, dan aksen biru tunggal. Tanpa emoji."""
    st.markdown(f"""
    <style>
    html, body, [class*="css"], .stApp, button, input, textarea, select, .stMarkdown {{
        font-family: {_APPLE_FONT_STACK} !important;
        color: {COLOR_TEXT_PRIMARY};
    }}

    .stApp {{
        background-color: {COLOR_BG};
    }}
    [data-testid="stAppViewContainer"], [data-testid="stHeader"] {{
        background-color: {COLOR_BG};
    }}
    .block-container {{
        padding-top: 2.2rem;
        padding-bottom: 3rem;
        max-width: 1280px;
    }}

    /* ---------- MACRO TYPOGRAPHY: Hero Header ---------- */
    .hero-wrap {{
        padding: 12px 4px 36px 4px;
    }}
    .hero-eyebrow {{
        color: {COLOR_TEXT_SECONDARY};
        font-size: 15px;
        font-weight: 600;
        letter-spacing: 0.01em;
        margin-bottom: 6px;
    }}
    .hero-title {{
        font-size: clamp(42px, 6vw, 84px);
        font-weight: 800;
        letter-spacing: -0.04em;
        line-height: 1.02;
        color: {COLOR_TEXT_PRIMARY};
        margin: 0;
    }}
    .hero-subtitle {{
        font-size: clamp(16px, 1.6vw, 21px);
        font-weight: 400;
        color: {COLOR_TEXT_SECONDARY};
        letter-spacing: -0.01em;
        margin-top: 14px;
        max-width: 760px;
        line-height: 1.45;
    }}

    .section-title {{
        font-size: 30px;
        font-weight: 800;
        letter-spacing: -0.03em;
        color: {COLOR_TEXT_PRIMARY};
        margin: 6px 0 22px 0;
    }}
    .section-subtitle {{
        font-size: 15px;
        color: {COLOR_TEXT_SECONDARY};
        margin-top: -14px;
        margin-bottom: 22px;
    }}

    /* ---------- BENTO BOX ---------- */
    .bento {{
        background: {COLOR_CARD};
        border-radius: 26px;
        border: 1px solid {COLOR_BORDER};
        box-shadow: 0 20px 40px rgba(0,0,0,0.04);
        padding: 28px 30px;
        margin-bottom: 26px;
        transition: transform 0.28s cubic-bezier(0.22,1,0.36,1), box-shadow 0.28s ease;
    }}
    .bento:hover {{
        transform: translateY(-4px);
        box-shadow: 0 28px 56px rgba(0,0,0,0.08);
    }}
    .bento-tight {{
        padding: 22px 26px;
    }}

    /* ---------- KPI Bento ---------- */
    .kpi-label {{
        font-size: 13px;
        font-weight: 600;
        color: {COLOR_TEXT_SECONDARY};
        text-transform: uppercase;
        letter-spacing: 0.04em;
        margin-bottom: 10px;
    }}
    .kpi-value {{
        font-size: 46px;
        font-weight: 800;
        letter-spacing: -0.03em;
        color: {COLOR_TEXT_PRIMARY};
        line-height: 1;
    }}
    .kpi-sub {{
        font-size: 13px;
        font-weight: 500;
        color: {COLOR_ACCENT};
        margin-top: 10px;
    }}

    /* ---------- Issue Card ---------- */
    .issue-title {{
        font-size: 19px;
        font-weight: 700;
        letter-spacing: -0.02em;
        color: {COLOR_TEXT_PRIMARY};
        margin-bottom: 6px;
    }}
    .issue-meta {{
        font-size: 12.5px;
        color: {COLOR_TEXT_SECONDARY};
        margin-bottom: 14px;
        line-height: 1.7;
    }}
    .badge {{
        display: inline-block;
        padding: 4px 13px;
        border-radius: 20px;
        font-size: 11.5px;
        font-weight: 700;
        margin-right: 8px;
        letter-spacing: 0.01em;
    }}
    .issue-section-label {{
        font-weight: 700;
        color: {COLOR_TEXT_SECONDARY};
        font-size: 11.5px;
        text-transform: uppercase;
        letter-spacing: 0.05em;
        margin-top: 14px;
    }}
    .issue-section-text {{
        font-size: 14.5px;
        color: {COLOR_TEXT_PRIMARY};
        line-height: 1.6;
        margin-top: 4px;
    }}

    /* ---------- Sidebar ---------- */
    section[data-testid="stSidebar"] {{
        background-color: {COLOR_CARD};
        border-right: 1px solid {COLOR_BORDER};
    }}
    section[data-testid="stSidebar"] * {{
        color: {COLOR_TEXT_PRIMARY} !important;
    }}
    section[data-testid="stSidebar"] label {{
        font-weight: 600 !important;
        font-size: 13px !important;
        color: {COLOR_TEXT_SECONDARY} !important;
    }}
    section[data-testid="stSidebar"] .stButton button {{
        background: {COLOR_ACCENT};
        color: white !important;
        border: none;
        font-weight: 600;
        border-radius: 12px;
        padding: 10px 0;
    }}
    section[data-testid="stSidebar"] .stButton button:hover {{
        background: {COLOR_ACCENT_DARK};
    }}
    .sidebar-title {{
        font-size: 22px;
        font-weight: 800;
        letter-spacing: -0.03em;
        margin-bottom: 4px;
    }}
    .sidebar-caption {{
        font-size: 12.5px;
        color: {COLOR_TEXT_SECONDARY};
        margin-bottom: 24px;
    }}

    /* ---------- Buttons ---------- */
    .stButton button {{
        border-radius: 12px;
        font-weight: 600;
        font-size: 14.5px;
        border: 1px solid rgba(0,0,0,0.08);
        background: {COLOR_CARD};
        color: {COLOR_TEXT_PRIMARY};
        padding: 10px 18px;
        transition: all 0.18s ease;
    }}
    .stButton button:hover {{
        border-color: {COLOR_ACCENT};
        color: {COLOR_ACCENT};
    }}
    div[data-testid="stFormSubmitButton"] button {{
        background: {COLOR_ACCENT};
        color: white;
        border: none;
        border-radius: 14px;
        padding: 13px 20px;
        font-size: 15.5px;
        font-weight: 700;
        letter-spacing: -0.01em;
    }}
    div[data-testid="stFormSubmitButton"] button:hover {{
        background: {COLOR_ACCENT_DARK};
        color: white;
    }}

    /* ---------- Input Fields ---------- */
    .stTextInput input, .stTextArea textarea, .stNumberInput input,
    .stDateInput input, div[data-baseweb="select"] > div {{
        border-radius: 12px !important;
        border: 1px solid rgba(0,0,0,0.08) !important;
        background: {COLOR_BG} !important;
        font-size: 14.5px !important;
    }}
    .stTextInput input:focus, .stTextArea textarea:focus {{
        border-color: {COLOR_ACCENT} !important;
        box-shadow: 0 0 0 4px rgba(0,113,227,0.12) !important;
    }}

    /* ---------- Tabs ---------- */
    button[data-baseweb="tab"] {{
        font-weight: 600;
        font-size: 15.5px;
        letter-spacing: -0.01em;
        color: {COLOR_TEXT_SECONDARY};
        padding: 10px 6px;
    }}
    button[data-baseweb="tab"][aria-selected="true"] {{
        color: {COLOR_TEXT_PRIMARY};
    }}
    div[data-baseweb="tab-highlight"] {{
        background-color: {COLOR_ACCENT} !important;
        height: 3px !important;
    }}
    div[data-baseweb="tab-border"] {{
        background-color: rgba(0,0,0,0.06) !important;
    }}

    /* ---------- Expander ---------- */
    div[data-testid="stExpander"] {{
        background: {COLOR_CARD};
        border: 1px solid {COLOR_BORDER};
        border-radius: 20px;
        box-shadow: 0 12px 28px rgba(0,0,0,0.03);
    }}
    .streamlit-expanderHeader {{
        font-weight: 650;
        font-size: 15px;
    }}

    mark {{
        background-color: {COLOR_HIGHLIGHT} !important;
        color: {COLOR_TEXT_PRIMARY} !important;
        padding: 0 3px;
        border-radius: 4px;
    }}

    hr {{
        border-color: rgba(0,0,0,0.06);
        margin: 28px 0;
    }}

    ::-webkit-scrollbar {{ width: 10px; height: 10px; }}
    ::-webkit-scrollbar-thumb {{ background: rgba(0,0,0,0.15); border-radius: 10px; }}
    </style>
    """, unsafe_allow_html=True)


def badge_html(label: str, colors: dict) -> str:
    """Membuat span badge status/dampak dengan warna pucat sesuai token."""
    safe_label = html_lib.escape(str(label))
    return (f'<span class="badge" style="background:{colors["bg"]};'
            f'color:{colors["text"]};">{safe_label}</span>')


# ==============================================================================
# 7. KOMPONEN UI: KPI BENTO CARDS
# ==============================================================================

def render_kpi_cards(df: pd.DataFrame):
    total_issues = len(df)
    total_verified = len(df[df["status"] == STATUS_VERIFIED]) if total_issues > 0 else 0
    pct_verified = (total_verified / total_issues * 100) if total_issues > 0 else 0.0
    total_high_impact = len(df[df["impact_level"] == "Tinggi"]) if total_issues > 0 else 0
    total_documents = len(df[df["file_name"].astype(str).str.strip() != ""]) if total_issues > 0 else 0

    col1, col2, col3, col4 = st.columns(4)
    kpi_data = [
        (col1, "Total Isu", f"{total_issues:,}", "Seluruh entri tercatat"),
        (col2, "Terverifikasi", f"{pct_verified:.1f}%", f"{total_verified} dari {total_issues} entri"),
        (col3, "Dampak Tinggi", f"{total_high_impact:,}", "Memerlukan prioritas"),
        (col4, "Dokumen Terunggah", f"{total_documents:,}", "File pendukung tersimpan"),
    ]
    for col, label, value, sub in kpi_data:
        with col:
            st.markdown(f"""
            <div class="bento bento-tight">
                <div class="kpi-label">{label}</div>
                <div class="kpi-value">{value}</div>
                <div class="kpi-sub">{sub}</div>
            </div>
            """, unsafe_allow_html=True)


# ==============================================================================
# 8. HALAMAN: DASHBOARD
# ==============================================================================

def page_dashboard(df: pd.DataFrame):
    st.markdown('<div class="section-title">Ringkasan Eksekutif</div>', unsafe_allow_html=True)
    render_kpi_cards(df)

    if df.empty:
        st.markdown("""
        <div class="bento">
            <div class="issue-section-text">Belum ada data isu. Unggah isu baru melalui tab
            <b>Unggah Baru</b> untuk mulai membangun basis pengetahuan strategis Anda.</div>
        </div>
        """, unsafe_allow_html=True)
        return

    col_a, col_b = st.columns([1.1, 1])
    with col_a:
        st.markdown('<div class="bento">', unsafe_allow_html=True)
        st.markdown("**Distribusi Status Governance**")
        status_counts = df["status"].value_counts().reset_index()
        status_counts.columns = ["Status", "Jumlah"]
        color_map = {s: STATUS_COLORS[s]["text"] for s in ALL_STATUSES}
        fig = px.pie(status_counts, names="Status", values="Jumlah", hole=0.62,
                     color="Status", color_discrete_map=color_map)
        fig.update_traces(textinfo="percent+label", marker=dict(line=dict(color=COLOR_CARD, width=3)))
        fig.update_layout(showlegend=False, margin=dict(t=10, b=10, l=10, r=10), height=310)
        st.plotly_chart(fig, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with col_b:
        st.markdown('<div class="bento">', unsafe_allow_html=True)
        st.markdown("**Isu Berdasarkan Level Dampak**")
        impact_counts = df["impact_level"].value_counts().reindex(IMPACT_LEVELS).fillna(0).reset_index()
        impact_counts.columns = ["Dampak", "Jumlah"]
        color_map_impact = {lvl: IMPACT_COLORS[lvl]["text"] for lvl in IMPACT_LEVELS}
        fig2 = px.bar(impact_counts, x="Dampak", y="Jumlah", color="Dampak",
                      color_discrete_map=color_map_impact, text="Jumlah")
        fig2.update_layout(showlegend=False, margin=dict(t=10, b=10, l=10, r=10), height=310,
                           xaxis_title=None, yaxis_title=None)
        fig2.update_traces(marker=dict(cornerradius=8))
        st.plotly_chart(fig2, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-title" style="font-size:24px;">Isu Terbaru</div>', unsafe_allow_html=True)
    recent = df.sort_values("id", ascending=False).head(5)
    for _, row in recent.iterrows():
        render_issue_card(row, search_keyword="")


# ==============================================================================
# 9. KOMPONEN UI: KARTU ISU
# ==============================================================================

def render_issue_card(row, search_keyword: str = ""):
    title_html = highlight_keyword(row["title"], search_keyword)
    summary_html = highlight_keyword(row["summary"], search_keyword)
    root_cause_html = highlight_keyword(row["root_cause"], search_keyword)
    recommendation_html = highlight_keyword(row["recommendation"], search_keyword)

    project_name = row["project_name"] if pd.notna(row["project_name"]) and row["project_name"] else "-"
    file_info = row["file_name"] if pd.notna(row["file_name"]) and row["file_name"] else "Tidak ada dokumen"

    status_badge = badge_html(row["status"], STATUS_COLORS.get(row["status"], STATUS_COLORS[STATUS_DRAFT]))
    impact_badge = badge_html(f'Dampak {row["impact_level"]}', IMPACT_COLORS.get(row["impact_level"], IMPACT_COLORS["Sedang"]))

    st.markdown(f"""
    <div class="bento">
        <div class="issue-title">{title_html}</div>
        <div class="issue-meta">
            {html_lib.escape(str(row['category']))} &nbsp;&middot;&nbsp;
            Proyek: {html_lib.escape(str(project_name))} &nbsp;&middot;&nbsp;
            {html_lib.escape(str(row['uploader']))} &nbsp;&middot;&nbsp;
            {html_lib.escape(str(row['upload_date']))} &nbsp;&middot;&nbsp;
            Versi {html_lib.escape(str(row['document_version']))} &nbsp;&middot;&nbsp;
            {html_lib.escape(str(file_info))}
        </div>
        {impact_badge}{status_badge}
        <div class="issue-section-label">Ringkasan Isu</div>
        <div class="issue-section-text">{summary_html if summary_html else '<i>Belum diisi</i>'}</div>
        <div class="issue-section-label">Akar Masalah</div>
        <div class="issue-section-text">{root_cause_html if root_cause_html else '<i>Belum diisi</i>'}</div>
        <div class="issue-section-label">Rekomendasi / Tindak Lanjut</div>
        <div class="issue-section-text">{recommendation_html if recommendation_html else '<i>Belum diisi</i>'}</div>
    </div>
    """, unsafe_allow_html=True)


# ==============================================================================
# 10. HALAMAN: TELUSURI (Pencarian & Filter)
# ==============================================================================

def page_search(df: pd.DataFrame, filters: dict):
    st.markdown('<div class="section-title">Telusuri Isu Strategis</div>', unsafe_allow_html=True)

    search_query = st.text_input(
        "Cari berdasarkan judul, ringkasan, akar masalah, atau rekomendasi",
        placeholder="Contoh: keterlambatan pengadaan, risiko anggaran",
        key="main_search_box"
    )

    filtered_df = apply_filters(df, filters)

    if search_query and search_query.strip():
        query_lower = search_query.strip().lower()
        mask = (
            filtered_df["title"].astype(str).str.lower().str.contains(re.escape(query_lower), na=False) |
            filtered_df["summary"].astype(str).str.lower().str.contains(re.escape(query_lower), na=False) |
            filtered_df["root_cause"].astype(str).str.lower().str.contains(re.escape(query_lower), na=False) |
            filtered_df["recommendation"].astype(str).str.lower().str.contains(re.escape(query_lower), na=False) |
            filtered_df["project_name"].astype(str).str.lower().str.contains(re.escape(query_lower), na=False)
        )
        filtered_df = filtered_df[mask]

    st.markdown(f'<div class="section-subtitle">{len(filtered_df)} isu ditemukan dari total {len(df)} entri</div>',
                unsafe_allow_html=True)

    if filtered_df.empty:
        st.markdown("""
        <div class="bento">
            <div class="issue-section-text">Tidak ada isu yang cocok dengan pencarian atau filter Anda.
            Coba ubah kata kunci atau kombinasi filter di sidebar.</div>
        </div>
        """, unsafe_allow_html=True)
        return

    sort_option = st.selectbox(
        "Urutkan berdasarkan", ["Terbaru", "Terlama", "Dampak Tertinggi", "Judul A sampai Z"], index=0
    )
    if sort_option == "Terbaru":
        filtered_df = filtered_df.sort_values("id", ascending=False)
    elif sort_option == "Terlama":
        filtered_df = filtered_df.sort_values("id", ascending=True)
    elif sort_option == "Dampak Tertinggi":
        impact_order = {"Tinggi": 0, "Sedang": 1, "Rendah": 2}
        filtered_df = filtered_df.assign(_ord=filtered_df["impact_level"].map(impact_order)).sort_values("_ord")
    elif sort_option == "Judul A sampai Z":
        filtered_df = filtered_df.sort_values("title", ascending=True)

    for _, row in filtered_df.iterrows():
        render_issue_card(row, search_keyword=search_query)


def apply_filters(df: pd.DataFrame, filters: dict) -> pd.DataFrame:
    if df.empty:
        return df
    result = df.copy()
    try:
        if filters.get("categories"):
            result = result[result["category"].isin(filters["categories"])]
        if filters.get("impacts"):
            result = result[result["impact_level"].isin(filters["impacts"])]
        if filters.get("statuses"):
            result = result[result["status"].isin(filters["statuses"])]
        if filters.get("year_range"):
            y_min, y_max = filters["year_range"]
            result = result[
                (result["project_year"].fillna(0).astype(int) >= y_min) &
                (result["project_year"].fillna(0).astype(int) <= y_max)
            ]
        return result
    except Exception as e:
        st.warning(f"Sebagian filter gagal diterapkan: {e}")
        return df


# ==============================================================================
# 11. HALAMAN: UNGGAH BARU (dengan Smart Auto-Fill)
# ==============================================================================

def page_upload():
    st.markdown('<div class="section-title">Unggah Isu Strategis Baru</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-subtitle">Unggah dokumen pendukung untuk mengaktifkan Smart Auto-Fill, '
                'atau isi form secara manual.</div>', unsafe_allow_html=True)

    st.markdown('<div class="bento">', unsafe_allow_html=True)
    st.markdown("**Langkah 1 — Dokumen Pendukung (Opsional)**")
    uploaded_file = st.file_uploader("Format didukung: PDF, DOCX, TXT", type=["pdf", "docx", "txt"], key="doc_uploader")

    if "autofill_result" not in st.session_state:
        st.session_state.autofill_result = {"summary": "", "root_cause": "", "recommendation": ""}
    if "extracted_text_cache" not in st.session_state:
        st.session_state.extracted_text_cache = ""

    if uploaded_file is not None:
        if st.button("Jalankan Smart Auto-Fill dari Dokumen Ini", use_container_width=True):
            with st.spinner("Membaca dan menganalisis dokumen..."):
                extracted_text = parse_uploaded_document(uploaded_file)
                st.session_state.extracted_text_cache = extracted_text
                if extracted_text.strip():
                    st.session_state.autofill_result = smart_auto_fill(extracted_text)
                    st.success("Berhasil mengekstrak dan menyarankan isi form di bawah. Tinjau dan sunting seperlunya.")
                else:
                    st.warning("Tidak ada teks yang berhasil diekstrak dari dokumen ini. Isi form secara manual.")
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="bento">', unsafe_allow_html=True)
    st.markdown("**Langkah 2 — Detail Isu**")

    with st.form("upload_issue_form", clear_on_submit=False):
        col1, col2 = st.columns(2)
        with col1:
            title = st.text_input("Judul Isu", placeholder="Contoh: Keterlambatan Pengadaan Material Utama")
            project_name = st.text_input("Nama Proyek", placeholder="Contoh: Proyek Pembangunan Gedung A")
            category = st.selectbox("Kategori Isu", CATEGORY_OPTIONS)
            impact_level = st.selectbox("Level Dampak", IMPACT_LEVELS, index=1)
        with col2:
            uploader = st.text_input("Nama Pengunggah / Pemilik Isu", placeholder="Contoh: Budi Santoso")
            project_year = st.number_input("Tahun Proyek", min_value=2000, max_value=2100,
                                            value=datetime.now().year, step=1)
            document_version = st.text_input("Versi Dokumen", value="v1.0")
            upload_date_input = st.date_input("Tanggal Diunggah", value=date.today())

        summary = st.text_area(
            "Ringkasan Isu", height=100,
            value=st.session_state.autofill_result.get("summary", ""),
            placeholder="Jelaskan secara singkat isu atau masalah yang terjadi"
        )
        root_cause = st.text_area(
            "Akar Masalah", height=100,
            value=st.session_state.autofill_result.get("root_cause", ""),
            placeholder="Jelaskan akar penyebab dari isu tersebut"
        )
        recommendation = st.text_area(
            "Rekomendasi Solusi / Tindak Lanjut", height=100,
            value=st.session_state.autofill_result.get("recommendation", ""),
            placeholder="Jelaskan rekomendasi atau tindakan korektif yang diusulkan"
        )

        submitted = st.form_submit_button("Simpan Isu (Status Draft)", use_container_width=True)

        if submitted:
            errors = []
            if not title or not title.strip():
                errors.append("Judul Isu wajib diisi.")
            if not uploader or not uploader.strip():
                errors.append("Nama Pengunggah wajib diisi.")
            if not summary or not summary.strip():
                errors.append("Ringkasan Isu wajib diisi.")
            if not root_cause or not root_cause.strip():
                errors.append("Akar Masalah wajib diisi.")
            if not recommendation or not recommendation.strip():
                errors.append("Rekomendasi Solusi wajib diisi.")

            if errors:
                for err in errors:
                    st.error(err)
            else:
                data = {
                    "title": title,
                    "project_name": project_name,
                    "category": category,
                    "impact_level": impact_level,
                    "status": STATUS_DRAFT,
                    "summary": summary,
                    "root_cause": root_cause,
                    "recommendation": recommendation,
                    "uploader": uploader,
                    "upload_date": upload_date_input.strftime("%Y-%m-%d"),
                    "project_year": int(project_year),
                    "document_version": document_version if document_version.strip() else "v1.0",
                    "file_name": uploaded_file.name if uploaded_file is not None else "",
                    "extracted_text": st.session_state.extracted_text_cache,
                }
                success = insert_issue(data)
                if success:
                    st.success("Isu berhasil disimpan dengan status Draft. Menunggu verifikasi PMO/Manager.")
                    st.session_state.autofill_result = {"summary": "", "root_cause": "", "recommendation": ""}
                    st.session_state.extracted_text_cache = ""
    st.markdown('</div>', unsafe_allow_html=True)


# ==============================================================================
# 12. HALAMAN: APPROVAL PMO (Governance Workflow)
# ==============================================================================

def page_approval_center(df: pd.DataFrame):
    st.markdown('<div class="section-title">Approval PMO</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-subtitle">Tinjau dan tetapkan status governance untuk setiap entri isu.</div>',
                unsafe_allow_html=True)

    pending_df = df[df["status"] == STATUS_DRAFT] if not df.empty else df

    tab_pending, tab_all = st.tabs([f"Menunggu Review ({len(pending_df)})", "Riwayat Seluruh Status"])

    with tab_pending:
        if pending_df.empty:
            st.markdown("""
            <div class="bento">
                <div class="issue-section-text">Tidak ada entri yang menunggu review saat ini.
                Seluruh isu sudah diverifikasi atau ditindaklanjuti.</div>
            </div>
            """, unsafe_allow_html=True)
        else:
            for _, row in pending_df.iterrows():
                with st.expander(f"{row['title']}  —  Diunggah oleh {row['uploader']} ({row['upload_date']})"):
                    st.markdown(f"**Kategori:** {row['category']} &nbsp;&middot;&nbsp; **Dampak:** {row['impact_level']} "
                                f"&nbsp;&middot;&nbsp; **Proyek:** {row['project_name'] or '-'} "
                                f"&nbsp;&middot;&nbsp; **Versi Dokumen:** {row['document_version']}")
                    st.markdown(f"**Ringkasan Isu**\n\n{row['summary']}")
                    st.markdown(f"**Akar Masalah**\n\n{row['root_cause']}")
                    st.markdown(f"**Rekomendasi**\n\n{row['recommendation']}")
                    if row['file_name']:
                        st.markdown(f"**Dokumen Pendukung:** {row['file_name']}")

                    reviewer_notes = st.text_area(
                        "Catatan Reviewer (opsional)", key=f"notes_{row['id']}",
                        placeholder="Tambahkan catatan verifikasi di sini"
                    )
                    col_a, col_b, col_c = st.columns(3)
                    with col_a:
                        if st.button("Verifikasi", key=f"approve_{row['id']}", use_container_width=True):
                            if update_issue_status(row['id'], STATUS_VERIFIED, reviewer_notes):
                                st.success(f"Isu '{row['title']}' berhasil diverifikasi.")
                                st.rerun()
                    with col_b:
                        if st.button("Tolak / Arsipkan", key=f"reject_{row['id']}", use_container_width=True):
                            if update_issue_status(row['id'], STATUS_REJECTED, reviewer_notes):
                                st.warning(f"Isu '{row['title']}' ditolak dan diarsipkan.")
                                st.rerun()
                    with col_c:
                        if st.button("Hapus Permanen", key=f"delete_{row['id']}", use_container_width=True):
                            if delete_issue(row['id']):
                                st.warning(f"Isu '{row['title']}' dihapus permanen dari database.")
                                st.rerun()

    with tab_all:
        if df.empty:
            st.markdown('<div class="bento"><div class="issue-section-text">Belum ada data untuk ditampilkan.</div></div>',
                        unsafe_allow_html=True)
        else:
            display_cols = ["id", "title", "category", "impact_level", "status",
                             "uploader", "upload_date", "reviewed_by", "reviewed_date", "reviewer_notes"]
            st.markdown('<div class="bento">', unsafe_allow_html=True)
            st.dataframe(df[display_cols], use_container_width=True, hide_index=True)
            st.markdown('</div>', unsafe_allow_html=True)


# ==============================================================================
# 13. DASHBOARD ANALITIK MENDALAM (disertakan dalam tab Dashboard sebagai bagian tambahan)
# ==============================================================================

def render_deep_analytics(df: pd.DataFrame):
    if df.empty:
        return

    st.markdown('<div class="section-title" style="font-size:24px;">Analitik Mendalam</div>', unsafe_allow_html=True)

    col1, col2 = st.columns([1.3, 1])
    with col1:
        st.markdown('<div class="bento">', unsafe_allow_html=True)
        st.markdown("**Tren Isu berdasarkan Kategori**")
        cat_counts = df["category"].value_counts().reset_index()
        cat_counts.columns = ["Kategori", "Jumlah"]
        fig = px.bar(cat_counts, x="Jumlah", y="Kategori", orientation="h",
                     color="Jumlah", color_continuous_scale=["#E5E5EA", COLOR_ACCENT])
        fig.update_layout(margin=dict(t=10, b=10, l=10, r=10), height=360,
                          yaxis=dict(categoryorder="total ascending"), coloraxis_showscale=False)
        fig.update_traces(marker=dict(cornerradius=8))
        st.plotly_chart(fig, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="bento">', unsafe_allow_html=True)
        st.markdown("**Tren Jumlah Isu per Tahun**")
        try:
            year_trend = df.groupby("project_year").size().reset_index(name="Jumlah").sort_values("project_year")
            fig3 = px.line(year_trend, x="project_year", y="Jumlah", markers=True)
            fig3.update_traces(line_color=COLOR_ACCENT, line_width=3,
                               marker=dict(size=9, color=COLOR_TEXT_PRIMARY))
            fig3.update_layout(margin=dict(t=10, b=10, l=10, r=10), height=360,
                               xaxis_title=None, yaxis_title=None)
            st.plotly_chart(fig3, use_container_width=True)
        except Exception as e:
            st.warning(f"Tidak dapat menampilkan grafik tren tahunan: {e}")
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="bento">', unsafe_allow_html=True)
    st.markdown("**Status Governance per Kategori**")
    cross = df.groupby(["category", "status"]).size().reset_index(name="Jumlah")
    color_map = {s: STATUS_COLORS[s]["text"] for s in ALL_STATUSES}
    fig2 = px.bar(cross, x="category", y="Jumlah", color="status", barmode="stack",
                  color_discrete_map=color_map)
    fig2.update_layout(margin=dict(t=10, b=10, l=10, r=10), height=340,
                       xaxis_title=None, yaxis_title=None, legend_title_text="Status")
    st.plotly_chart(fig2, use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)


# ==============================================================================
# 14. EXPORT DATA (disertakan di tab Dashboard sebagai bagian tambahan)
# ==============================================================================

def render_export_section(df: pd.DataFrame, filters: dict):
    st.markdown('<div class="section-title" style="font-size:24px;">Export Data</div>', unsafe_allow_html=True)
    st.markdown('<div class="bento">', unsafe_allow_html=True)
    export_scope = st.radio("Cakupan data", ["Seluruh Data", "Data Terfilter"], horizontal=True)
    export_df = apply_filters(df, filters) if export_scope == "Data Terfilter" else df
    st.markdown(f"Total baris yang akan diekspor: **{len(export_df)}**")
    st.dataframe(export_df, use_container_width=True, hide_index=True)

    col1, col2 = st.columns(2)
    with col1:
        try:
            csv_bytes = export_df.to_csv(index=False).encode("utf-8-sig")
            st.download_button("Unduh sebagai CSV", data=csv_bytes,
                               file_name=f"lessons_learned_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                               mime="text/csv", use_container_width=True)
        except Exception as e:
            st.error(f"Gagal membuat file CSV: {e}")
    with col2:
        try:
            excel_bytes = dataframe_to_excel_bytes(export_df)
            if excel_bytes:
                st.download_button("Unduh sebagai Excel", data=excel_bytes,
                                   file_name=f"lessons_learned_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                   use_container_width=True)
        except Exception as e:
            st.error(f"Gagal membuat file Excel: {e}")
    st.markdown('</div>', unsafe_allow_html=True)


# ==============================================================================
# 15. SIDEBAR: FILTER GLOBAL & NAVIGASI PERAN
# ==============================================================================

def render_sidebar(df: pd.DataFrame) -> dict:
    with st.sidebar:
        st.markdown('<div class="sidebar-title">Filter</div>', unsafe_allow_html=True)
        st.markdown('<div class="sidebar-caption">Navigasi dan penyaringan data</div>', unsafe_allow_html=True)

        role = st.selectbox("Login sebagai", ["Staff / Uploader", "PMO / Manager"], index=0)
        st.session_state["current_role"] = role

        st.markdown("---")

        categories_available = sorted(df["category"].dropna().unique().tolist()) if not df.empty else CATEGORY_OPTIONS
        selected_categories = st.multiselect("Kategori Isu", categories_available, default=[])
        selected_impacts = st.multiselect("Level Dampak", IMPACT_LEVELS, default=[])
        selected_statuses = st.multiselect("Status Verifikasi", ALL_STATUSES, default=[])

        if not df.empty and df["project_year"].notna().any():
            min_year = int(df["project_year"].min())
            max_year = int(df["project_year"].max())
        else:
            min_year, max_year = 2015, datetime.now().year
        if min_year == max_year:
            max_year = min_year + 1
        year_range = st.slider("Rentang Tahun Proyek", min_value=min_year, max_value=max_year,
                                value=(min_year, max_year))

        st.markdown("---")
        if st.button("Reset Semua Filter", use_container_width=True):
            st.rerun()

        st.markdown("---")
        st.caption("Enterprise Knowledge Management Platform — 2026")

    return {
        "categories": selected_categories,
        "impacts": selected_impacts,
        "statuses": selected_statuses,
        "year_range": year_range,
    }


# ==============================================================================
# 16. FUNGSI UTAMA (MAIN)
# ==============================================================================

def main():
    st.set_page_config(
        page_title="Enterprise Knowledge Management",
        page_icon=None,
        layout="wide",
        initial_sidebar_state="expanded"
    )

    init_db()
    load_custom_css()

    st.markdown("""
    <div class="hero-wrap">
        <div class="hero-eyebrow">Enterprise Knowledge Management</div>
        <div class="hero-title">Strategic Lessons Learned.</div>
        <div class="hero-subtitle">
            Dokumentasikan, telusuri, dan analisis isu strategis proyek —
            dari akar masalah hingga tindakan korektif, dalam satu platform terpusat.
        </div>
    </div>
    """, unsafe_allow_html=True)

    df = fetch_all_issues()
    filters = render_sidebar(df)

    tabs = st.tabs(["Dashboard", "Telusuri", "Unggah Baru", "Approval PMO"])

    with tabs[0]:
        try:
            filtered_for_dashboard = apply_filters(df, filters)
            page_dashboard(filtered_for_dashboard)
            st.markdown("<hr/>", unsafe_allow_html=True)
            render_deep_analytics(filtered_for_dashboard)
            st.markdown("<hr/>", unsafe_allow_html=True)
            render_export_section(df, filters)
        except Exception as e:
            st.error(f"Terjadi kesalahan pada Dashboard: {e}")

    with tabs[1]:
        try:
            page_search(df, filters)
        except Exception as e:
            st.error(f"Terjadi kesalahan pada halaman Telusuri: {e}")

    with tabs[2]:
        try:
            page_upload()
        except Exception as e:
            st.error(f"Terjadi kesalahan pada halaman Unggah: {e}")

    with tabs[3]:
        try:
            if st.session_state.get("current_role") == "PMO / Manager":
                page_approval_center(df)
            else:
                st.markdown("""
                <div class="bento">
                    <div class="issue-section-text">Halaman ini khusus untuk peran <b>PMO / Manager</b>.
                    Ubah peran Anda di sidebar untuk mengakses Approval PMO.</div>
                </div>
                """, unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Terjadi kesalahan pada Approval PMO: {e}")


# ==============================================================================
# ENTRY POINT
# ==============================================================================

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error("Terjadi kesalahan tak terduga pada aplikasi.")
        st.exception(e)
